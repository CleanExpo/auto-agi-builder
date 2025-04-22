import os
import subprocess
import json
import logging
import time
from pathlib import Path
from typing import Dict, List, Optional, Tuple

class AutonomousBuilder:
    """
    Autonomous build system for the Prototype Generation Platform.
    Configured specifically for Windows environment.
    Only requires manual input for environment variables.
    """
    
    def __init__(self, project_root: str = None):
        if project_root is None:
            # Default to the current location you specified
            project_root = r"C:\Users\PhillMcGurk\OneDrive - Disaster Recovery\1111\Auto AGI Builder"
        self.project_root = Path(project_root).absolute()
        self.env_file = self.project_root / ".env"
        self.required_env_vars = [
            "OPENAI_API_KEY",
            "FIREBASE_API_KEY",
            "FIREBASE_AUTH_DOMAIN",
            "FIREBASE_PROJECT_ID",
            "POSTGRES_USER",
            "POSTGRES_PASSWORD",
            "POSTGRES_HOST",
            "POSTGRES_DB",
            "MONGO_URI",
            "REDIS_URL"
        ]
        
        # Setup logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(self.project_root / "build.log"),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger("autonomous_builder")
        
    def check_environment(self) -> bool:
        """Check if all required tools are installed (Windows version)"""
        required_tools = {
            "python": "3.10",
            "pip": "22.0",
            "node": "18.0",
            "npm": "8.0",
            "docker": "24.0"
        }
        
        self.logger.info("Checking required tools...")
        
        for tool, min_version in required_tools.items():
            try:
                if tool in ["python", "pip"]:
                    process = subprocess.run([tool, "--version"], capture_output=True, text=True, shell=True)
                elif tool == "docker":
                    process = subprocess.run(["docker", "--version"], capture_output=True, text=True, shell=True)
                else:
                    process = subprocess.run([tool, "-v"], capture_output=True, text=True, shell=True)
                
                if process.returncode != 0:
                    self.logger.error(f"{tool} not found")
                    return False
                
                self.logger.info(f"{tool} found: {process.stdout.strip()}")
            except FileNotFoundError:
                self.logger.error(f"{tool} not found")
                return False
        
        return True
    
    def check_env_variables(self) -> bool:
        """Check if all required environment variables are set"""
        if not self.env_file.exists():
            self.logger.warning(".env file not found. Creating template...")
            self._create_env_template()
            self.logger.info(f"Please fill in the .env file at {self.env_file}")
            return False
        
        env_vars = {}
        with open(self.env_file, "r") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#"):
                    key, value = line.split("=", 1)
                    env_vars[key] = value
        
        missing_vars = [var for var in self.required_env_vars if var not in env_vars or not env_vars[var]]
        
        if missing_vars:
            self.logger.error(f"Missing environment variables: {', '.join(missing_vars)}")
            return False
        
        self.logger.info("All required environment variables found")
        return True
    
    def _create_env_template(self):
        """Create a template .env file"""
        with open(self.env_file, "w") as f:
            f.write("# Required environment variables for the Prototype Generation Platform\n\n")
            for var in self.required_env_vars:
                f.write(f"{var}=\n")
    
    def setup_project_structure(self):
        """Create the project directory structure (Windows compatible)"""
        self.logger.info("Setting up project directory structure...")
        
        directories = [
            "app\\api",
            "app\\core",
            "app\\db",
            "app\\services",
            "app\\ai",
            "app\\utils",
            "frontend",
            "tests"
        ]
        
        for directory in directories:
            dir_path = self.project_root / directory
            dir_path.mkdir(parents=True, exist_ok=True)
            self.logger.info(f"Created directory: {dir_path}")
        
        # Create __init__.py files
        for directory in directories:
            if directory.startswith("app"):
                init_file = self.project_root / directory / "__init__.py"
                init_file.touch(exist_ok=True)
        
        self.logger.info("Project directory structure created successfully")
    
    # All other methods remain the same as in the previous version
    # ...

    def run_windows_installer(self):
        """Run the Windows-specific installation steps"""
        self.logger.info("Starting Windows installation process...")
        
        # Create a batch file to run the setup
        batch_file = self.project_root / "setup.bat"
        with open(batch_file, "w") as f:
            f.write("""@echo off
echo Setting up the Prototype Generation Platform...

echo Checking Python installation...
python --version
if %ERRORLEVEL% neq 0 (
    echo Python not found. Please install Python 3.10 or later.
    exit /b 1
)

echo Installing Python dependencies...
pip install -r requirements.txt
if %ERRORLEVEL% neq 0 (
    echo Failed to install Python dependencies.
    exit /b 1
)

echo Setting up frontend...
cd frontend
npm install
if %ERRORLEVEL% neq 0 (
    echo Failed to install frontend dependencies.
    exit /b 1
)
cd ..

echo Setup complete!
echo To start the application, run: python -m app.main
""")
        
        self.logger.info(f"Created Windows setup batch file at {batch_file}")
        self.logger.info("To complete setup, run: setup.bat")
        
        # Create a run script
        run_file = self.project_root / "run.bat"
        with open(run_file, "w") as f:
            f.write("""@echo off
echo Starting the Prototype Generation Platform...

start cmd /k "cd frontend && npm run dev"
start cmd /k "python -m app.main"

echo Application started!
echo Frontend: http://localhost:3000
echo API: http://localhost:8000
""")
        
        self.logger.info(f"Created Windows run batch file at {run_file}")
        self.logger.info("To start the application, run: run.bat")
        
        return True

    def generate_requirements_txt(self):
        """Generate the requirements.txt file"""
        self.logger.info("Generating requirements.txt...")
        requirements_content = """\
fastapi
uvicorn[standard]
python-dotenv
psycopg2-binary
pymongo
redis
openai
google-cloud-firestore # Assuming Firebase interaction might use Firestore
requests
"""
        requirements_file = self.project_root / "requirements.txt"
        with open(requirements_file, "w") as f:
            f.write(requirements_content)
        self.logger.info(f"Generated requirements.txt at {requirements_file}")

    def generate_frontend_foundation(self):
        """Generate the basic Next.js frontend structure."""
        self.logger.info("Generating Next.js frontend foundation...")
        frontend_dir = self.project_root / "frontend"
        frontend_dir.mkdir(parents=True, exist_ok=True) # Ensure frontend dir exists

        # --- Update package.json ---
        package_json_file = frontend_dir / "package.json"
        package_json_content = {
            "name": "prototype-frontend",
            "version": "0.1.0",
            "private": True,
            "dependencies": {
                "react": "^18", # Use major version for broader compatibility
                "react-dom": "^18",
                "firebase": "^10", # Use major version
                "next": "^14" # Add Next.js
            },
            "scripts": {
                "dev": "next dev", # Next.js dev script (runs on port 3000 by default)
                "build": "next build",
                "start": "next start", # Next.js production start
                "lint": "next lint"
            },
            "eslintConfig": {
                "extends": [
                    "next/core-web-vitals" # Use Next.js specific linting
                ]
            },
            "browserslist": {
                "production": [
                    ">0.2%",
                    "not dead",
                    "not op_mini all"
                ],
                "development": [
                    "last 1 chrome version",
                    "last 1 firefox version",
                    "last 1 safari version"
                ]
            }
        }
        with open(package_json_file, "w") as f:
            json.dump(package_json_content, f, indent=2)
        self.logger.info(f"Updated {package_json_file} for Next.js")

        # --- Create directories ---
        dirs_to_create = ["pages", "components", "public", "styles"]
        for d in dirs_to_create:
            dir_path = frontend_dir / d
            dir_path.mkdir(exist_ok=True)
            self.logger.info(f"Ensured directory exists: {dir_path}")
            # Add .gitkeep to components dir to ensure it's tracked if empty initially
            if d == "components":
                (dir_path / ".gitkeep").touch(exist_ok=True)

        # --- Create basic files ---

        # next.config.js
        next_config_file = frontend_dir / "next.config.js"
        next_config_content = """\
/** @type {import('next').NextConfig} */
const nextConfig = {
  reactStrictMode: true,
}

module.exports = nextConfig
"""
        with open(next_config_file, "w") as f:
            f.write(next_config_content)
        self.logger.info(f"Generated {next_config_file}")

        # styles/globals.css
        globals_css_file = frontend_dir / "styles" / "globals.css"
        globals_css_content = """\
/* Basic Reset */
html,
body {
  padding: 0;
  margin: 0;
  font-family: -apple-system, BlinkMacSystemFont, Segoe UI, Roboto, Oxygen,
    Ubuntu, Cantarell, Fira Sans, Droid Sans, Helvetica Neue, sans-serif;
}

a {
  color: inherit;
  text-decoration: none;
}

* {
  box-sizing: border-box;
}

/* Add some basic styling */
body {
  padding: 2rem;
}
"""
        with open(globals_css_file, "w") as f:
            f.write(globals_css_content)
        self.logger.info(f"Generated {globals_css_file}")

        # pages/_app.js
        app_js_file = frontend_dir / "pages" / "_app.js"
        app_js_content = """\
import '../styles/globals.css'

function MyApp({ Component, pageProps }) {
  return <Component {...pageProps} />
}

export default MyApp
"""
        with open(app_js_file, "w") as f:
            f.write(app_js_content)
        self.logger.info(f"Generated {app_js_file}")

        # styles/Home.module.css
        home_css_file = frontend_dir / "styles" / "Home.module.css"
        home_css_content = """\
/* Example CSS Module */
.container {
  min-height: 100vh;
  padding: 0 0.5rem;
  display: flex;
  flex-direction: column;
  justify-content: center;
  align-items: center;
}

.main {
  padding: 5rem 0;
  flex: 1;
  display: flex;
  flex-direction: column;
  justify-content: center;
  align-items: center;
}

.title {
  margin: 0;
  line-height: 1.15;
  font-size: 4rem;
  text-align: center;
}

.description {
  line-height: 1.5;
  font-size: 1.5rem;
  text-align: center;
}
"""
        with open(home_css_file, "w") as f:
             f.write(home_css_content)
        self.logger.info(f"Generated {home_css_file}")

        # --- Create Components ---

        # components/MeetingInputForm.js
        input_form_file = frontend_dir / "components" / "MeetingInputForm.js"
        input_form_content = """\
import React, { useState } from 'react';

// Basic styling (can be moved to CSS modules later)
const styles = {
  form: {
    display: 'flex',
    flexDirection: 'column',
    gap: '1rem',
    width: '100%',
    maxWidth: '600px',
    margin: '2rem auto',
  },
  textarea: {
    minHeight: '150px',
    padding: '0.5rem',
    fontSize: '1rem',
    fontFamily: 'inherit',
    border: '1px solid #ccc',
    borderRadius: '4px',
  },
  button: {
    padding: '0.75rem 1.5rem',
    fontSize: '1rem',
    cursor: 'pointer',
    backgroundColor: '#0070f3',
    color: 'white',
    border: 'none',
    borderRadius: '4px',
  },
  buttonDisabled: {
    backgroundColor: '#ccc',
    cursor: 'not-allowed',
  },
  error: {
    color: 'red',
    marginTop: '0.5rem',
  },
  loading: {
    marginTop: '0.5rem',
    fontStyle: 'italic',
    textAlign: 'center',
  }
};

export default function MeetingInputForm({ onSubmit, isLoading, error }) {
  const [notes, setNotes] = useState('');

  const handleSubmit = (event) => {
    event.preventDefault();
    if (!notes.trim()) {
      alert('Please enter some meeting notes.');
      return;
    }
    onSubmit(notes);
  };

  return (
    <form onSubmit={handleSubmit} style={styles.form}>
      <label htmlFor="meetingNotes">Enter Meeting Notes:</label>
      <textarea
        id="meetingNotes"
        value={notes}
        onChange={(e) => setNotes(e.target.value)}
        placeholder="Paste or type your meeting notes here..."
        style={styles.textarea}
        disabled={isLoading}
      />
      <button
        type="submit"
        style={{...styles.button, ...(isLoading ? styles.buttonDisabled : {})}}
        disabled={isLoading}
      >
        {isLoading ? 'Generating...' : 'Generate Prototype Analysis'}
      </button>
      {isLoading && <p style={styles.loading}>Processing, please wait...</p>}
      {error && <p style={styles.error}>Error: {error}</p>}
    </form>
  );
}
"""
        with open(input_form_file, "w") as f:
            f.write(input_form_content)
        self.logger.info(f"Generated {input_form_file}")

        # components/PrototypeDisplay.js
        display_file = frontend_dir / "components" / "PrototypeDisplay.js"
        display_content = """\
import React from 'react';

// Basic styling
const styles = {
  container: {
    marginTop: '2rem',
    padding: '1rem',
    border: '1px solid #ccc',
    borderRadius: '4px',
    width: '100%',
    maxWidth: '800px',
    margin: '2rem auto',
    fontFamily: 'monospace',
    backgroundColor: '#f9f9f9',
  },
  heading: {
    marginTop: '1rem',
    marginBottom: '0.5rem',
    borderBottom: '1px solid #eee',
    paddingBottom: '0.3rem',
    fontSize: '1.2em',
  },
  pre: {
    whiteSpace: 'pre-wrap', // Preserve formatting
    wordWrap: 'break-word',
    maxHeight: '300px', // Limit height for long responses
    overflowY: 'auto', // Add scroll for long responses
    backgroundColor: '#fff',
    padding: '0.5rem',
    border: '1px solid #eee',
    borderRadius: '4px',
  },
  error: {
    color: 'red',
    fontWeight: 'bold',
  },
  errorDetails: {
     color: '#555',
     fontSize: '0.9em',
  }
};

// Helper to render sections nicely
const renderSection = (title, data) => {
  let content;
  let error = null;
  let rawResponse = null;

  if (!data) {
    content = <p>No data available for this section.</p>;
  } else if (typeof data === 'object' && data.error) {
    error = data.error;
    rawResponse = data.raw_response;
    content = <p style={styles.error}>{error}</p>;
    if (rawResponse) {
        content = (
            <>
                {content}
                <p style={styles.errorDetails}>Raw Response:</p>
                <pre style={styles.pre}>{rawResponse}</pre>
            </>
        );
    }
  } else {
    // Simple JSON stringify for display
    try {
        content = <pre style={styles.pre}>{JSON.stringify(data, null, 2)}</pre>;
    } catch (e) {
        content = <p style={styles.error}>Error displaying data.</p>;
    }
  }

  return (
    <div>
      <h3 style={styles.heading}>{title}</h3>
      {content}
    </div>
  );
};


export default function PrototypeDisplay({ analysisResult }) {
  if (!analysisResult) {
    return null; // Don't render anything if there's no result yet
  }

  // Check for top-level API error first
  if (analysisResult.detail) {
      return (
          <div style={{...styles.container, borderColor: 'red'}}>
              <h2>Error</h2>
              <p style={styles.error}>{analysisResult.detail}</p>
          </div>
      );
  }

  return (
    <div style={styles.container}>
      <h2>Prototype Analysis Results</h2>
      {renderSection("Requirements", analysisResult.requirements)}
      {renderSection("Prototype Details", analysisResult.prototype_details)}
      {renderSection("Integrations", analysisResult.integrations)}
      {renderSection("Pricing", analysisResult.pricing)}
    </div>
  );
}
"""
        with open(display_file, "w") as f:
            f.write(display_content)
        self.logger.info(f"Generated {display_file}")


        # pages/index.js - Update to use components
        index_js_file = frontend_dir / "pages" / "index.js"
        index_js_content = """\
import Head from 'next/head'
import styles from '../styles/Home.module.css' // Now this file exists
import React, { useState } from 'react'; // Import useState
import MeetingInputForm from '../components/MeetingInputForm'; // Import component
import PrototypeDisplay from '../components/PrototypeDisplay'; // Import component

export default function Home() {
  const [analysisResult, setAnalysisResult] = useState(null);
  const [isLoading, setIsLoading] = useState(false);
  const [error, setError] = useState(null);

  const handleFormSubmit = async (notes) => {
    setIsLoading(true);
    setError(null);
    setAnalysisResult(null); // Clear previous results

    try {
      // Use the correct API endpoint URL
      const apiUrl = 'http://localhost:8000/api/v1/prototype/generate'; // Backend runs on 8000
      console.log(`Sending request to ${apiUrl}`);
      const response = await fetch(apiUrl, {
        method: 'POST',
        headers: {
          'Content-Type': 'application/json',
          'Accept': 'application/json', // Explicitly accept JSON
        },
        body: JSON.stringify({ notes: notes }),
      });

      console.log(`Received response status: ${response.status}`);

      let result;
      try {
          result = await response.json();
          console.log("API Response Body:", result);
      } catch (parseError) {
          console.error("Failed to parse JSON response:", parseError);
          // If JSON parsing fails, maybe the response wasn't JSON (e.g., HTML error page)
          throw new Error(`HTTP error! status: ${response.status}. Failed to parse response.`);
      }


      if (!response.ok) {
        console.error("API Error Response:", result);
        // Use the detail field from FastAPI's HTTPException or a default message
        throw new Error(result.detail || `HTTP error! status: ${response.status}`);
      }

      setAnalysisResult(result);

    } catch (err) {
      console.error("Failed to fetch analysis:", err);
      setError(err.message || 'Failed to fetch analysis.');
      setAnalysisResult(null); // Ensure no partial results are shown on error
    } finally {
      setIsLoading(false);
    }
  };


  return (
    <div className={styles.container}>
      <Head>
        <title>Auto AGI Builder</title>
        <meta name="description" content="Prototype Generation Platform" />
        <link rel="icon" href="/favicon.ico" /> {/* Assumes favicon in public dir */}
      </Head>

      <main className={styles.main}>
        <h1 className={styles.title}>
          Auto AGI Builder
        </h1>

        <p className={styles.description}>
          Enter meeting notes below to generate a prototype analysis.
        </p>

        {/* Add the form component */}
        <MeetingInputForm
          onSubmit={handleFormSubmit}
          isLoading={isLoading}
          error={error}
        />

        {/* Add the display component - conditionally render */}
        {analysisResult && <PrototypeDisplay analysisResult={analysisResult} />}
        {/* Display error message if fetch failed */}
        {!isLoading && error && <p style={{color: 'red', marginTop: '1rem'}}>API Error: {error}</p>}

      </main>

      <footer>
        {/* Footer content */}
      </footer>
    </div>
  )
}
"""
        with open(index_js_file, "w") as f:
            f.write(index_js_content)
        self.logger.info(f"Generated {index_js_file}")

        self.logger.info("Next.js frontend foundation generated successfully.")

    def generate_backend_foundation(self):
        """Generate the core backend files (main.py, config.py)"""
        self.logger.info("Generating backend foundation...")

        # --- Create app/core/config.py ---
        core_dir = self.project_root / "app" / "core"
        core_dir.mkdir(parents=True, exist_ok=True)
        config_file = core_dir / "config.py"
        config_content = """\
# app/core/config.py
import os
from pydantic_settings import BaseSettings
from dotenv import load_dotenv
from pathlib import Path

# Load .env file from the project root (adjust path as necessary)
# Assuming .env is in the parent directory of 'app'
env_path = Path(__file__).resolve().parent.parent.parent / '.env'
load_dotenv(dotenv_path=env_path)

class Settings(BaseSettings):
    PROJECT_NAME: str = "Auto AGI Builder"
    API_V1_STR: str = "/api/v1"

    # Environment variables loaded from .env
    OPENAI_API_KEY: str = os.getenv("OPENAI_API_KEY", "")
    FIREBASE_API_KEY: str = os.getenv("FIREBASE_API_KEY", "")
    FIREBASE_AUTH_DOMAIN: str = os.getenv("FIREBASE_AUTH_DOMAIN", "")
    FIREBASE_PROJECT_ID: str = os.getenv("FIREBASE_PROJECT_ID", "")
    POSTGRES_USER: str = os.getenv("POSTGRES_USER", "")
    POSTGRES_PASSWORD: str = os.getenv("POSTGRES_PASSWORD", "")
    POSTGRES_HOST: str = os.getenv("POSTGRES_HOST", "")
    POSTGRES_DB: str = os.getenv("POSTGRES_DB", "")
    MONGO_URI: str = os.getenv("MONGO_URI", "")
    REDIS_URL: str = os.getenv("REDIS_URL", "")

    # CORS settings - Allow frontend origin
    BACKEND_CORS_ORIGINS: list[str] = ["http://localhost:3000"] # Default frontend dev port

    class Config:
        case_sensitive = True

settings = Settings()
"""
        with open(config_file, "w") as f:
            f.write(config_content)
        self.logger.info(f"Generated {config_file}")

        # --- Create app/main.py ---
        app_dir = self.project_root / "app"
        main_file = app_dir / "main.py"
        main_content = """\
# app/main.py
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
import logging
import os
import sys

# Ensure app modules can be imported
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from core.config import settings
except ImportError as e:
    print(f"Error importing settings: {e}")
    print(f"Current sys.path: {sys.path}")
    # Fallback settings if import fails during initial setup/debug
    class FallbackSettings:
        PROJECT_NAME = "Auto AGI Builder (Fallback)"
        API_V1_STR = "/api/v1"
        BACKEND_CORS_ORIGINS = ["http://localhost:3000"]
    settings = FallbackSettings()


# Setup basic logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title=settings.PROJECT_NAME,
    openapi_url=f"{settings.API_V1_STR}/openapi.json"
)

# Set CORS middleware
if settings.BACKEND_CORS_ORIGINS:
    logger.info(f"Adding CORS middleware for origins: {settings.BACKEND_CORS_ORIGINS}")
    app.add_middleware(
        CORSMiddleware,
        allow_origins=[str(origin).strip() for origin in settings.BACKEND_CORS_ORIGINS],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )
else:
    logger.warning("No CORS origins specified.")

# Placeholder root endpoint
@app.get("/", tags=["Root"])
async def read_root():
    logger.info("Root endpoint accessed")
    return {"message": f"Welcome to {settings.PROJECT_NAME}"}

# Import and include the main API router
try:
    from api.v1.api import api_router
    app.include_router(api_router, prefix=settings.API_V1_STR)
    logger.info(f"API router included at prefix: {settings.API_V1_STR}")
except ImportError as e:
    logger.error(f"Could not import or include api_router: {e}")

# Note: The run.bat starts this using 'python -m app.main'
# This means uvicorn needs to be callable programmatically or via CLI.
# The __main__ block below is useful for direct `python app/main.py` execution,
# but might not be hit when run via `python -m`.
# Uvicorn is typically run from the command line pointing to the app object.
# Example: uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload

# If run directly using `python app/main.py`
if __name__ == "__main__":
    import uvicorn
    logger.info("Attempting to start Uvicorn server directly...")
    # This might fail if run via `python -m app.main` as __name__ won't be "__main__"
    # uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True, app_dir="app") # Use CLI instead for reliability
    print("To run the server with reload, use: uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload")


"""
        with open(main_file, "w") as f:
            f.write(main_content)
        self.logger.info(f"Generated {main_file}")

        # Create __init__.py in app/core if it doesn't exist
        core_init = core_dir / "__init__.py"
        core_init.touch(exist_ok=True)
        # Create __init__.py in app if it doesn't exist (should already be there)
        app_init = app_dir / "__init__.py"
        app_init.touch(exist_ok=True)

        self.logger.info("Backend foundation generated successfully.")

    def generate_ai_components(self):
        """Generate the AI component stubs"""
        self.logger.info("Generating AI components...")
        ai_dir = self.project_root / "app" / "ai"
        ai_dir.mkdir(parents=True, exist_ok=True)

        # Create __init__.py
        init_file = ai_dir / "__init__.py"
        init_content = "# app/ai/__init__.py\n"
        with open(init_file, "w") as f:
            f.write(init_content)
        self.logger.info(f"Generated {init_file}")

        # Create llm_service.py
        llm_service_file = ai_dir / "llm_service.py"
        llm_service_content = """\
# app/ai/llm_service.py
import os
from openai import OpenAI
import logging
# Need to ensure core.config is importable from here
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
try:
    from core.config import settings
except ImportError:
    # Fallback if run standalone or path issues
    class FallbackSettings: OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    settings = FallbackSettings()


logger = logging.getLogger(__name__)

class LLMService:
    def __init__(self):
        if not settings.OPENAI_API_KEY:
            logger.error("OPENAI_API_KEY not found in settings.")
            # Avoid raising error during setup, just log
            self.client = None
            return
            # raise ValueError("OpenAI API Key not configured.")
        try:
            self.client = OpenAI(api_key=settings.OPENAI_API_KEY)
            logger.info("OpenAI client initialized successfully.")
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI client: {e}")
            self.client = None
            # raise # Avoid raising during setup

    def generate_text(self, prompt: str, model: str = "gpt-4", max_tokens: int = 1500) -> str:
        \"\"\"Generates text using the configured OpenAI model.\"\"\"
        if not self.client:
            logger.error("OpenAI client not initialized. Cannot generate text.")
            return "Error: OpenAI client not available."

        logger.info(f"Generating text with model {model} for prompt: {prompt[:50]}...")
        try:
            response = self.client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
            )
            if response.choices and len(response.choices) > 0:
                if response.choices[0].message and response.choices[0].message.content:
                    generated_text = response.choices[0].message.content.strip()
                    logger.info("Text generation successful.")
                    return generated_text
                else:
                    logger.error("No message content found in OpenAI response choice.")
                    return "Error: No content in response."
            else:
                logger.error("No choices found in OpenAI response.")
                return "Error: No choices in response."
        except Exception as e:
            logger.error(f"Error during OpenAI API call: {e}")
            return f"Error generating text: {e}"

# Singleton instance
try:
    llm_service = LLMService()
except Exception: # Catch potential init errors during setup
    llm_service = None
    logger.error("LLMService singleton failed to initialize.")

"""
        with open(llm_service_file, "w") as f:
            f.write(llm_service_content)
        self.logger.info(f"Generated {llm_service_file}")

        # Create requirements_extractor.py
        req_extractor_file = ai_dir / "requirements_extractor.py"
        req_extractor_content = """\
# app/ai/requirements_extractor.py
import logging
import json
from .llm_service import llm_service # Use the singleton instance

logger = logging.getLogger(__name__)

class RequirementsExtractor:
    def extract(self, meeting_notes: str) -> dict:
        \"\"\"
        Extracts structured requirements from meeting notes using the LLM.
        (MCP 6: meeting-notes-capture -> pain-point-extraction -> feature-identification)
        \"\"\"
        if not llm_service or not llm_service.client:
             logger.error("LLM service not available for requirements extraction.")
             return {"error": "LLM service not available."}

        logger.info("Extracting requirements from meeting notes...")
        prompt = f'''
Analyze the following meeting notes and extract the key requirements, pain points, and desired features.
Structure the output as a JSON object with keys: "pain_points" (list of strings), "features" (list of strings), "constraints" (list of strings).

Meeting Notes:
---
{meeting_notes}
---

JSON Output:
'''
        try:
            response_text = llm_service.generate_text(prompt)
            try:
                # Attempt to find JSON block if LLM adds extra text
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                if json_start != -1 and json_end != -1:
                     response_text = response_text[json_start:json_end]

                extracted_data = json.loads(response_text)
                # Basic validation
                if isinstance(extracted_data.get("pain_points"), list) and \
                   isinstance(extracted_data.get("features"), list) and \
                   isinstance(extracted_data.get("constraints"), list):
                    logger.info("Requirements extracted successfully.")
                    return extracted_data
                else:
                    logger.warning(f"Extracted JSON has incorrect structure: {response_text}")
                    return {"error": "Extracted JSON structure invalid", "raw_response": response_text}
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse LLM response as JSON: {e}\\nResponse: {response_text}")
                return {"error": "Failed to parse requirements JSON", "raw_response": response_text}
        except Exception as e:
            logger.error(f"Error during requirements extraction: {e}")
            return {"error": str(e)}

requirements_extractor = RequirementsExtractor()
"""
        with open(req_extractor_file, "w") as f:
            f.write(req_extractor_content)
        self.logger.info(f"Generated {req_extractor_file}")

        # Create prototype_generator.py
        proto_gen_file = ai_dir / "prototype_generator.py"
        proto_gen_content = """\
# app/ai/prototype_generator.py
import logging
import json
from .llm_service import llm_service

logger = logging.getLogger(__name__)

class PrototypeGenerator:
    def generate(self, requirements: dict) -> dict:
        \"\"\"
        Generates prototype details based on requirements.
        (MCP 5: requirements-analysis -> component-selection -> layout-generation)
        \"\"\"
        if not llm_service or not llm_service.client:
             logger.error("LLM service not available for prototype generation.")
             return {"error": "LLM service not available."}

        logger.info("Generating prototype details...")
        requirements_str = json.dumps(requirements, indent=2)

        prompt = f'''
Based on the following requirements, suggest a basic prototype structure, key UI components,
and potential technologies. Output as a JSON object with keys: "suggested_structure" (string description),
"ui_components" (list of strings), "technologies" (list of strings).

Requirements:
---
{requirements_str}
---

JSON Output:
'''
        try:
            response_text = llm_service.generate_text(prompt)
            try:
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                if json_start != -1 and json_end != -1:
                     response_text = response_text[json_start:json_end]

                prototype_details = json.loads(response_text)
                if isinstance(prototype_details.get("suggested_structure"), str) and \
                   isinstance(prototype_details.get("ui_components"), list) and \
                   isinstance(prototype_details.get("technologies"), list):
                    logger.info("Prototype details generated successfully.")
                    return prototype_details
                else:
                    logger.warning(f"Prototype details JSON has incorrect structure: {response_text}")
                    return {"error": "Prototype details JSON structure invalid", "raw_response": response_text}
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse LLM response as JSON: {e}\\nResponse: {response_text}")
                return {"error": "Failed to parse prototype details JSON", "raw_response": response_text}
        except Exception as e:
            logger.error(f"Error during prototype generation: {e}")
            return {"error": str(e)}

prototype_generator = PrototypeGenerator()
"""
        with open(proto_gen_file, "w") as f:
            f.write(proto_gen_content)
        self.logger.info(f"Generated {proto_gen_file}")

        # Create integration_mapper.py
        int_mapper_file = ai_dir / "integration_mapper.py"
        int_mapper_content = """\
# app/ai/integration_mapper.py
import logging
import json
from .llm_service import llm_service

logger = logging.getLogger(__name__)

class IntegrationMapper:
    def map(self, requirements: dict) -> dict:
        \"\"\"
        Suggests potential service integrations based on requirements.
        (MCP 2: service-integrations)
        \"\"\"
        if not llm_service or not llm_service.client:
             logger.error("LLM service not available for integration mapping.")
             return {"error": "LLM service not available."}

        logger.info("Mapping potential integrations...")
        requirements_str = json.dumps(requirements, indent=2)

        prompt = f'''
Based on the following requirements, suggest potential third-party services or APIs
that could be integrated (e.g., payment gateways, CRM, analytics).
Output as a JSON object with keys: "suggested_integrations" (list of strings with brief description for each).

Requirements:
---
{requirements_str}
---

JSON Output:
'''
        try:
            response_text = llm_service.generate_text(prompt)
            try:
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                if json_start != -1 and json_end != -1:
                     response_text = response_text[json_start:json_end]

                integrations = json.loads(response_text)
                if isinstance(integrations.get("suggested_integrations"), list):
                    logger.info("Integrations mapped successfully.")
                    return integrations
                else:
                     logger.warning(f"Integrations JSON has incorrect structure: {response_text}")
                     return {"error": "Integrations JSON structure invalid", "raw_response": response_text}
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse LLM response as JSON: {e}\\nResponse: {response_text}")
                return {"error": "Failed to parse integrations JSON", "raw_response": response_text}
        except Exception as e:
            logger.error(f"Error during integration mapping: {e}")
            return {"error": str(e)}

integration_mapper = IntegrationMapper()
"""
        with open(int_mapper_file, "w") as f:
            f.write(int_mapper_content)
        self.logger.info(f"Generated {int_mapper_file}")

        # Create pricing_engine.py
        pricing_file = ai_dir / "pricing_engine.py"
        pricing_content = """\
# app/ai/pricing_engine.py
import logging
import json
from .llm_service import llm_service

logger = logging.getLogger(__name__)

class PricingEngine:
    def calculate(self, requirements: dict, prototype_details: dict) -> dict:
        \"\"\"
        Calculates an estimated price based on requirements and prototype complexity.
        (MCP 3: pricing-engine, MCP 6: value-proposition-calculation, MCP 7: pricing-presentation)
        \"\"\"
        if not llm_service or not llm_service.client:
             logger.error("LLM service not available for pricing calculation.")
             return {"error": "LLM service not available."}

        logger.info("Calculating estimated price...")
        requirements_str = json.dumps(requirements, indent=2)
        prototype_str = json.dumps(prototype_details, indent=2)

        prompt = f'''
Based on the following requirements and suggested prototype details, provide a rough
estimated price range (e.g., low, high expressed as numbers) and a brief justification.
Output as a JSON object with keys: "estimated_price_low" (number), "estimated_price_high" (number), "justification" (string).

Requirements:
---
{requirements_str}
---

Prototype Details:
---
{prototype_str}
---

JSON Output:
'''
        try:
            response_text = llm_service.generate_text(prompt)
            try:
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                if json_start != -1 and json_end != -1:
                     response_text = response_text[json_start:json_end]

                pricing = json.loads(response_text)
                if isinstance(pricing.get("estimated_price_low"), (int, float)) and \
                   isinstance(pricing.get("estimated_price_high"), (int, float)) and \
                   isinstance(pricing.get("justification"), str):
                    logger.info("Pricing calculated successfully.")
                    return pricing
                else:
                    logger.error(f"Pricing JSON missing required keys or has wrong types. Response: {response_text}")
                    return {"error": "Pricing JSON structure/types invalid", "raw_response": response_text}
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse LLM response as JSON: {e}\\nResponse: {response_text}")
                return {"error": "Failed to parse pricing JSON", "raw_response": response_text}
        except Exception as e:
            logger.error(f"Error during pricing calculation: {e}")
            return {"error": str(e)}

pricing_engine = PricingEngine()
"""
        with open(pricing_file, "w") as f:
            f.write(pricing_content)
        self.logger.info(f"Generated {pricing_file}")

        self.logger.info("AI components generated successfully.")

    def integrate_document_import(self):
        """Integrate document import system with the main application"""
        self.logger.info("Integrating document import system...")
        
        # Update the main.py file to include the document import router
        main_py = self.project_root / "app" / "main.py"
        with open(main_py, "r") as f:
            content = f.read()
        
        # Add import statement if not already present
        if "from app.api import document_import" not in content:
            import_position = content.find("from fastapi import FastAPI")
            if import_position != -1:
                import_end = content.find("\n", import_position) + 1
                content = content[:import_end] + "from app.api import document_import\n" + content[import_end:]
        
        # Add router inclusion if not already present
        if "app.include_router(document_import.router)" not in content:
            app_declaration_pos = content.find("app = FastAPI")
            if app_declaration_pos != -1:
                # Find position after middleware setup
                middleware_pos = content.find("app.add_middleware", app_declaration_pos)
                if middleware_pos != -1:
                    middleware_end = content.find("\n\n", middleware_pos)
                    if middleware_end != -1:
                        content = content[:middleware_end+2] + "app.include_router(document_import.router)\n\n" + content[middleware_end+2:]
        
        # Write updated content back
        with open(main_py, "w") as f:
            f.write(content)
        
        # Update requirements.txt to include document processing libraries
        requirements = self.project_root / "requirements.txt"
        with open(requirements, "a") as f:
            f.write("\n# Document processing libraries\nPyPDF2==3.0.1\npython-docx==0.8.11\nopenpyxl==3.1.2\nbeautifulsoup4==4.12.2\nmarkdown==3.4.3\npyyaml==6.0.1\n")
        
        # Create a dedicated page for document import in the frontend
        documents_page = self.project_root / "frontend" / "pages" / "documents.js"
        with open(documents_page, "w") as f:
            f.write("""import { useState } from 'react';
import Head from 'next/head';
import DocumentImport from '../components/DocumentImport';
import DocumentAnalyzer from '../components/DocumentAnalyzer';

export default function Documents() {
  const [importedDocument, setImportedDocument] = useState(null);
  const [extractedRequirements, setExtractedRequirements] = useState(null);
  
  const handleDocumentImported = (document) => {
    setImportedDocument(document);
    // Reset requirements when a new document is imported
    setExtractedRequirements(null);
  };
  
  const handleRequirementsExtracted = (requirements) => {
    setExtractedRequirements(requirements);
    // Here you would typically update the main application state
    // with the extracted requirements
    console.log('Requirements extracted:', requirements);
  };
  
  return (
    <div>
      <Head>
        <title>Document Import - Prototype Generation Platform</title>
        <meta name="description" content="Import and analyze documents" />
      </Head>

      <main className="min-h-screen py-8 px-4 sm:px-6 lg:px-8">
        <div className="max-w-7xl mx-auto">
          <h1 className="text-3xl font-bold mb-8">Document Import</h1>
          
          <div className="grid grid-cols-1 md:grid-cols-2 gap-8">
            <div>
              <DocumentImport onDocumentImported={handleDocumentImported} />
            </div>
            <div>
              {importedDocument && (
                <DocumentAnalyzer 
                  document={importedDocument}
                  onRequirementsExtracted={handleRequirementsExtracted}
                />
              )}
            </div>
          </div>
          
          {extractedRequirements && (
            <div className="mt-8 bg-gray-50 p-4 rounded-lg">
              <h2 className="text-xl font-semibold mb-4">Extracted Requirements</h2>
              <pre className="bg-white p-4 rounded border overflow-auto max-h-96">
                {JSON.stringify(extractedRequirements, null, 2)}
              </pre>
            </div>
          )}
        </div>
      </main>
    </div>
  );
}
""")
        
        # Create the document import component
        document_import_component = self.project_root / "frontend" / "components" / "DocumentImport.js"
        with open(document_import_component, "w") as f:
            f.write("""import { useState } from 'react';

export default function DocumentImport({ onDocumentImported }) {
  const [file, setFile] = useState(null);
  const [uploading, setUploading] = useState(false);
  const [error, setError] = useState(null);
  
  const handleFileChange = (e) => {
    const selectedFile = e.target.files[0];
    if (selectedFile) {
      setFile(selectedFile);
      setError(null);
    }
  };
  
  const handleSubmit = async (e) => {
    e.preventDefault();
    
    if (!file) {
      setError("Please select a file to upload");
      return;
    }
    
    setUploading(true);
    setError(null);
    
    try {
      const formData = new FormData();
      formData.append('file', file);
      
      const response = await fetch('http://localhost:8000/api/document/upload', {
        method: 'POST',
        body: formData,
      });
      
      if (!response.ok) {
        throw new Error(`Upload failed: ${response.status}`);
      }
      
      const data = await response.json();
      setUploading(false);
      
      // Call the callback with the document info
      onDocumentImported({
        id: data.id,
        name: file.name,
        type: file.type,
        size: file.size,
        content: data.content,
      });
      
    } catch (err) {
      setError(err.message || 'Upload failed');
      setUploading(false);
    }
  };
  
  return (
    <div className="bg-white p-6 rounded-lg shadow-md">
      <h2 className="text-xl font-semibold mb-4">Import Document</h2>
      
      <form onSubmit={handleSubmit}>
        <div className="mb-4">
          <label className="block text-sm font-medium text-gray-700 mb-2">
            Select document
          </label>
          
          <input
            type="file"
            onChange={handleFileChange}
            className="w-full px-3 py-2 border border-gray-300 rounded-md"
            accept=".pdf,.docx,.doc,.txt,.md,.html,.json,.yaml,.yml"
          />
          
          <p className="mt-1 text-sm text-gray-500">
            Supported formats: PDF, Word, Text, Markdown, HTML, JSON, YAML
          </p>
        </div>
        
        {error && (
          <div className="mb-4 text-red-600 text-sm">
            {error}
          </div>
        )}
        
        <button
          type="submit"
          disabled={uploading || !file}
          className={`px-4 py-2 rounded-md text-white font-medium ${
            uploading || !file
              ? 'bg-blue-300 cursor-not-allowed'
              : 'bg-blue-600 hover:bg-blue-700'
          }`}
        >
          {uploading ? 'Uploading...' : 'Upload Document'}
        </button>
      </form>
    </div>
  );
}
""")
        
        # Create the document analyzer component
        document_analyzer = self.project_root / "frontend" / "components" / "DocumentAnalyzer.js"
        with open(document_analyzer, "w") as f:
            f.write("""import { useState } from 'react';

export default function DocumentAnalyzer({ document, onRequirementsExtracted }) {
  const [analyzing, setAnalyzing] = useState(false);
  const [error, setError] = useState(null);
  
  const analyzeDocument = async () => {
    if (!document) return;
    
    setAnalyzing(true);
    setError(null);
    
    try {
      const response = await fetch('http://localhost:8000/api/document/analyze', {
        method: 'POST',
        headers: {
          'Content-Type': 'application/json',
        },
        body: JSON.stringify({
          document_id: document.id,
        }),
      });
      
      if (!response.ok) {
        throw new Error(`Analysis failed: ${response.status}`);
      }
      
      const data = await response.json();
      setAnalyzing(false);
      
      // Call the callback with the extracted requirements
      onRequirementsExtracted(data);
      
    } catch (err) {
      setError(err.message || 'Analysis failed');
      setAnalyzing(false);
    }
  };
  
  if (!document) {
    return null;
  }
  
  return (
    <div className="bg-white p-6 rounded-lg shadow-md">
      <h2 className="text-xl font-semibold mb-4">Document Analyzer</h2>
      
      <div className="mb-4">
        <h3 className="font-medium text-gray-700">Document details:</h3>
        <p><strong>Name:</strong> {document.name}</p>
        <p><strong>Type:</strong> {document.type}</p>
        <p><strong>Size:</strong> {Math.round(document.size / 1024)} KB</p>
      </div>
      
      <div className="mb-4">
        <h3 className="font-medium text-gray-700 mb-2">Document content preview:</h3>
        <div className="bg-gray-50 p-3 rounded border overflow-auto max-h-60">
          <pre className="text-sm whitespace-pre-wrap">
            {document.content.slice(0, 500)}
            {document.content.length > 500 && '...'}
          </pre>
        </div>
      </div>
      
      {error && (
        <div className="mb-4 text-red-600 text-sm">
          {error}
        </div>
      )}
      
      <button
        onClick={analyzeDocument}
        disabled={analyzing}
        className={`px-4 py-2 rounded-md text-white font-medium ${
          analyzing
            ? 'bg-blue-300 cursor-not-allowed'
            : 'bg-blue-600 hover:bg-blue-700'
        }`}
      >
        {analyzing ? 'Analyzing...' : 'Extract Requirements'}
      </button>
    </div>
  );
}
""")
        
        # Create the API router for document import
        api_dir = self.project_root / "app" / "api"
        api_dir.mkdir(parents=True, exist_ok=True)
        
        document_import_router = api_dir / "document_import.py"
        with open(document_import_router, "w") as f:
            f.write("""import os
import uuid
import tempfile
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, HTTPException, Body
from fastapi.responses import JSONResponse
from typing import Dict, Any, Optional, List

# For different document types
import PyPDF2
try:
    import docx
    import openpyxl
    from bs4 import BeautifulSoup
    import markdown
    import yaml
except ImportError:
    # Handle import errors gracefully during setup
    pass

# Setup logging
import logging
logger = logging.getLogger(__name__)

# Create router
router = APIRouter(prefix="/api/document", tags=["Document Import"])

# In-memory storage for uploaded documents (for demo purposes)
# In a real application, this would be stored in a database
DOCUMENTS = {}

@router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    try:
        # Generate a unique document ID
        document_id = str(uuid.uuid4())
        
        # Create a temporary file to save the uploaded content
        suffix = Path(file.filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
            # Read the uploaded file in chunks
            content = await file.read()
            temp.write(content)
            temp_path = temp.name
        
        # Extract text from the document based on file type
        try:
            document_text = extract_text_from_file(temp_path, suffix)
        finally:
            # Clean up the temporary file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        
        # Store document info in memory
        DOCUMENTS[document_id] = {
            "filename": file.filename,
            "content_type": file.content_type,
            "text_content": document_text
        }
        
        return {
            "id": document_id,
            "content": document_text[:1000]  # Return first 1000 chars as preview
        }
        
    except Exception as e:
        logger.exception(f"Error processing document upload: {e}")
        raise HTTPException(status_code=500, detail=f"Document upload failed: {str(e)}")

@router.post("/analyze")
async def analyze_document(data: Dict[str, Any] = Body(...)):
    document_id = data.get("document_id")
    
    if not document_id or document_id not in DOCUMENTS:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # Get document content
        document = DOCUMENTS[document_id]
        
        # In a real application, this would call an AI service to analyze the document
        # For demo purposes, we'll generate a simple structure with extracted requirements
        
        # Split document into sections based on common patterns
        content = document["text_content"]
        
        # Very simple extraction logic (would be replaced with AI in production)
        requirements = extract_requirements_from_text(content)
        
        return requirements
        
    except Exception as e:
        logger.exception(f"Error analyzing document: {e}")
        raise HTTPException(status_code=500, detail=f"Document analysis failed: {str(e)}")

def extract_text_from_file(file_path: str, suffix: str) -> str:
    """Extract text content from different file types"""
    suffix = suffix.lower()
    
    try:
        # PDF files
        if suffix == ".pdf":
            return extract_text_from_pdf(file_path)
        
        # Word documents
        elif suffix in [".docx", ".doc"]:
            return extract_text_from_docx(file_path)
        
        # Plain text, markdown
        elif suffix in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            # Convert markdown to plain text if needed
            if suffix == ".md":
                try:
                    html = markdown.markdown(content)
                    soup = BeautifulSoup(html, "html.parser")
                    content = soup.get_text()
                except:
                    # Fall back to raw content if markdown conversion fails
                    pass
                    
            return content
        
        # HTML files
        elif suffix in [".html", ".htm"]:
            with open(file_path, "r", encoding="utf-8") as f:
                html = f.read()
            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text()
        
        # JSON files
        elif suffix == ".json":
            import json
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return json.dumps(data, indent=2)
        
        # YAML files
        elif suffix in [".yaml", ".yml"]:
            with open(file_path, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f)
            return yaml.dump(data, default_flow_style=False)
        
        # Default fallback - try to read as text
        else:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            except UnicodeDecodeError:
                return f"[Binary file type not supported: {suffix}]"
    
    except Exception as e:
        logger.exception(f"Error extracting text from {suffix} file: {e}")
        return f"[Error extracting text: {str(e)}]"

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = []
    
    try:
        with open(file_path, "rb") as f:
            pdf = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf.pages)):
                page = pdf.pages[page_num]
                text.append(page.extract_text())
        
        return "\\n".join(text)
    except Exception as e:
        logger.exception(f"Error extracting text from PDF: {e}")
        return f"[Error reading PDF: {str(e)}]"

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs]
        return "\\n".join(text)
    except Exception as e:
        logger.exception(f"Error extracting text from DOCX: {e}")
        return f"[Error reading DOCX: {str(e)}]"

def extract_requirements_from_text(text: str) -> Dict[str, Any]:
    """
    Extract requirements from document text.
    In a real application, this would use a sophisticated AI model.
    """
    # Simple keyword-based extraction for demo purposes
    lines = text.split("\\n")
    
    # Extract requirements based on common patterns
    requirements = []
    functional_requirements = []
    non_functional_requirements = []
    user_stories = []
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines
        if not line:
            continue
            
        # Look for requirement-like patterns
        lower_line = line.lower()
        
        # Simple pattern matching
        if any(keyword in lower_line for keyword in ["must", "should", "will", "shall"]):
            requirements.append(line)
            
            # Categorize by keywords
            if any(keyword in lower_line for keyword in ["performance", "security", "usability", "reliability"]):
                non_functional_requirements.append(line)
            else:
                functional_requirements.append(line)
                
        # Look for user story format: "As a ... I want ... So that ..."
        if "as a " in lower_line and " i want " in lower_line:
            user_stories.append(line)
    
    # Combine results
    return {
        "requirements": requirements,
        "functional_requirements": functional_requirements,
        "non_functional_requirements": non_functional_requirements,
        "user_stories": user_stories,
        "analysis_summary": f"Found {len(requirements)} requirements: {len(functional_requirements)} functional, {len(non_functional_requirements)} non-functional, and {len(user_stories)} user stories."
    }
""")
        
        # Add link to the main navigation in the frontend
        # We'll look for _app.js to add the link there
        app_js_file = self.project_root / "frontend" / "pages" / "_app.js"
        if app_js_file.exists():
            with open(app_js_file, "r") as f:
                app_js_content = f.read()
            
            # Check if there's navigation we can update
            if "function MyApp" in app_js_content and "</header>" in app_js_content:
                # Add a basic navigation if not present
                if "<nav" not in app_js_content:
                    updated_content = app_js_content.replace(
                        "function MyApp({ Component, pageProps }) {",
                        """function MyApp({ Component, pageProps }) {
  const Navigation = () => (
    <header className="bg-blue-600 text-white p-4">
      <div className="container mx-auto flex justify-between items-center">
        <a href="/" className="text-xl font-bold">Auto AGI Builder</a>
        <nav>
          <ul className="flex space-x-4">
            <li><a href="/" className="hover:underline">Home</a></li>
            <li><a href="/documents" className="hover:underline">Documents</a></li>
          </ul>
        </nav>
      </div>
    </header>
  );
"""
                    )
                    
                    # Also replace the return statement to include the navigation
                    updated_content = updated_content.replace(
                        "return <Component {...pageProps} />",
                        """return (
    <>
      <Navigation />
      <Component {...pageProps} />
    </>
  )"""
                    )
                    
                    with open(app_js_file, "w") as f:
                        f.write(updated_content)
        
        self.logger.info("Document import system integrated successfully.")

    def generate_api_endpoints(self):
        """Generate the API endpoint structure and basic prototype endpoint."""
        self.logger.info("Generating API endpoints...")
        api_dir = self.project_root / "app" / "api"
        v1_dir = api_dir / "v1"
        endpoints_dir = v1_dir / "endpoints"

        # Create directories if they don't exist
        endpoints_dir.mkdir(parents=True, exist_ok=True)

        # --- Create __init__.py files ---
        (api_dir / "__init__.py").touch(exist_ok=True)
        (v1_dir / "__init__.py").touch(exist_ok=True)
        endpoints_init_file = endpoints_dir / "__init__.py"
        endpoints_init_content = """\
# app/api/v1/endpoints/__init__.py
from . import prototype
"""
        with open(endpoints_init_file, "w") as f:
            f.write(endpoints_init_content)
        self.logger.info(f"Generated {endpoints_init_file}")

        # --- Create app/api/v1/api.py (Router aggregator) ---
        api_router_file = v1_dir / "api.py"
        api_router_content = """\
# app/api/v1/api.py
from fastapi import APIRouter
from .endpoints import prototype # Import endpoint modules

api_router = APIRouter()
api_router.include_router(prototype.router, prefix="/prototype", tags=["Prototype Generation"])

# Add other routers here if needed in the future
# e.g., api_router.include_router(users.router, prefix="/users", tags=["Users"])
"""
        with open(api_router_file, "w") as f:
            f.write(api_router_content)
        self.logger.info(f"Generated {api_router_file}")

        # --- Create app/api/v1/endpoints/prototype.py ---
        prototype_endpoint_file = endpoints_dir / "prototype.py"
        prototype_endpoint_content = """\
# app/api/v1/endpoints/prototype.py
from fastapi import APIRouter, HTTPException, Body
from pydantic import BaseModel, Field
import logging
import os
import sys
from typing import List, Dict, Any

# Ensure app modules can be imported
# Adjust path based on where the script is run from
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))))

# Attempt to import AI components, with fallbacks for setup phase
try:
    from app.ai.requirements_extractor import requirements_extractor
    from app.ai.prototype_generator import prototype_generator
    from app.ai.integration_mapper import integration_mapper
    from app.ai.pricing_engine import pricing_engine
    AI_LOADED = True
except ImportError as e:
    print(f"WARNING: Could not import AI components in prototype endpoint: {e}. Using dummies.")
    AI_LOADED = False
    class DummyExtractor: def extract(self, notes): return {"error": "Extractor not loaded", "pain_points": [], "features": [], "constraints": []}
    class DummyGenerator: def generate(self, reqs): return {"error": "Generator not loaded", "suggested_structure": "", "ui_components": [], "technologies": []}
    class DummyMapper: def map(self, reqs): return {"error": "Mapper not loaded", "suggested_integrations": []}
    class DummyPricer: def calculate(self, reqs, details): return {"error": "Pricer not loaded", "estimated_price_low": 0, "estimated_price_high": 0, "justification": ""}
    requirements_extractor = DummyExtractor()
    prototype_generator = DummyGenerator()
    integration_mapper = DummyMapper()
    pricing_engine = DummyPricer()


logger = logging.getLogger(__name__)
router = APIRouter()

# --- Input/Output Models ---
class MeetingNotesInput(BaseModel):
    notes: str = Field(..., description="Raw text meeting notes.")

class RequirementsOutput(BaseModel):
    pain_points: List[str] = Field(default_factory=list)
    features: List[str] = Field(default_factory=list)
    constraints: List[str] = Field(default_factory=list)
    error: Optional[str] = None
    raw_response: Optional[str] = None

class PrototypeDetailsOutput(BaseModel):
    suggested_structure: str = ""
    ui_components: List[str] = Field(default_factory=list)
    technologies: List[str] = Field(default_factory=list)
    error: Optional[str] = None
    raw_response: Optional[str] = None

class IntegrationsOutput(BaseModel):
    suggested_integrations: List[str] = Field(default_factory=list)
    error: Optional[str] = None
    raw_response: Optional[str] = None

class PricingOutput(BaseModel):
    estimated_price_low: float = 0.0
    estimated_price_high: float = 0.0
    justification: str = ""
    error: Optional[str] = None
    raw_response: Optional[str] = None

class FullPrototypeOutput(BaseModel):
    requirements: RequirementsOutput
    prototype_details: PrototypeDetailsOutput
    integrations: IntegrationsOutput
    pricing: PricingOutput

# --- API Endpoint ---
@router.post("/generate", response_model=FullPrototypeOutput, tags=["Prototype Generation"])
async def generate_full_prototype(payload: MeetingNotesInput = Body(...)):
    \"\"\"
    Takes meeting notes and generates a full prototype analysis including
    requirements, prototype details, integrations, and pricing using AI components.
    (Corresponds to MCPs 6, 5, 2, 3, 7)
    \"\"\"
    logger.info(f"Received request to generate prototype. AI Loaded: {AI_LOADED}")
    if not AI_LOADED:
         # If imports failed during startup, return dummy data with errors
         return FullPrototypeOutput(
             requirements=RequirementsOutput(**requirements_extractor.extract("")),
             prototype_details=PrototypeDetailsOutput(**prototype_generator.generate({})),
             integrations=IntegrationsOutput(**integration_mapper.map({})),
             pricing=PricingOutput(**pricing_engine.calculate({}, {}))
         )

    try:
        # MCP 6: Extract requirements
        logger.info("Step 1: Extracting requirements...")
        requirements_result = requirements_extractor.extract(payload.notes)
        if "error" in requirements_result:
            logger.error(f"Error during requirements extraction: {requirements_result}")
            # Return partial success with error indicated
            return FullPrototypeOutput(
                requirements=RequirementsOutput(**requirements_result),
                prototype_details=PrototypeDetailsOutput(error="Skipped due to requirements error"),
                integrations=IntegrationsOutput(error="Skipped due to requirements error"),
                pricing=PricingOutput(error="Skipped due to requirements error")
            )
        requirements_output = RequirementsOutput(**requirements_result) # Validate structure

        # MCP 5 & 3 (part 1): Generate prototype details
        logger.info("Step 2: Generating prototype details...")
        prototype_details_result = prototype_generator.generate(requirements_result)
        if "error" in prototype_details_result:
            logger.error(f"Error during prototype generation: {prototype_details_result}")
            return FullPrototypeOutput(
                requirements=requirements_output,
                prototype_details=PrototypeDetailsOutput(**prototype_details_result),
                integrations=IntegrationsOutput(error="Skipped due to prototype details error"),
                pricing=PricingOutput(error="Skipped due to prototype details error")
            )
        prototype_details_output = PrototypeDetailsOutput(**prototype_details_result)

        # MCP 2: Map integrations
        logger.info("Step 3: Mapping integrations...")
        integrations_result = integration_mapper.map(requirements_result)
        if "error" in integrations_result:
             logger.error(f"Error during integration mapping: {integrations_result}")
             # Continue to pricing if possible, but report error
             integrations_output = IntegrationsOutput(**integrations_result)
        else:
             integrations_output = IntegrationsOutput(**integrations_result)


        # MCP 3 (part 2) & 7: Calculate pricing
        logger.info("Step 4: Calculating pricing...")
        pricing_result = pricing_engine.calculate(requirements_result, prototype_details_result)
        if "error" in pricing_result:
            logger.error(f"Error during pricing calculation: {pricing_result}")
            pricing_output = PricingOutput(**pricing_result)
        else:
            pricing_output = PricingOutput(**pricing_result)


        logger.info("Prototype generation process completed.")
        return FullPrototypeOutput(
            requirements=requirements_output,
            prototype_details=prototype_details_output,
            integrations=integrations_output,
            pricing=pricing_output
        )

    except Exception as e:
        logger.exception("Unhandled exception during prototype generation process.")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")

"""
        with open(prototype_endpoint_file, "w") as f:
            f.write(prototype_endpoint_content)
        self.logger.info(f"Generated {prototype_endpoint_file}")

        self.logger.info("API endpoints generated successfully.")

    
    def execute_full_setup(self):
        """Execute the full setup process for Windows"""
        self.logger.info("Starting full setup process...")
        
        if not self.check_environment():
            self.logger.error("Environment check failed. Please install the required tools.")
            return False
        
        if not self.check_env_variables():
            self.logger.warning("Please fill in the environment variables before continuing.")
            return False
        
        self.setup_project_structure()
        self.generate_backend_foundation() # Now calls the actual method
        self.generate_frontend_foundation() # Now calls the actual method
        self.generate_requirements_txt() # Now calls the actual method
        # self.generate_dockerfile() # Placeholder
        self.generate_ai_components() # Now calls the actual method
        self.generate_api_endpoints() # Now calls the actual method
        
        self.run_windows_installer()
        
        self.logger.info("Setup process completed successfully!")
        self.logger.info("Next steps:")
        self.logger.info("1. Fill in the .env file with your API keys and configuration")
        self.logger.info("2. Run setup.bat to install dependencies")
        self.logger.info("3. Run run.bat to start the application")
        
        return True

# Add remaining methods from previous version...
# Placeholder for missing methods
def placeholder_method(self, *args, **kwargs):
    self.logger.info(f"Executing placeholder method: {args[0] if args else kwargs}")

# AutonomousBuilder.generate_backend_foundation = lambda self: placeholder_method(self, "generate_backend_foundation") # Removed placeholder
# AutonomousBuilder.generate_frontend_foundation = lambda self: placeholder_method(self, "generate_frontend_foundation") # Removed placeholder
# AutonomousBuilder.generate_requirements_txt = lambda self: placeholder_method(self, "generate_requirements_txt") # Removed placeholder
AutonomousBuilder.generate_dockerfile = lambda self: placeholder_method(self, "generate_dockerfile")
# AutonomousBuilder.generate_ai_components = lambda self: placeholder_method(self, "generate_ai_components") # Removed placeholder
# AutonomousBuilder.generate_api_endpoints = lambda self: placeholder_method(self, "generate_api_endpoints") # Removed placeholder


# Add a simple command-line interface
if __name__ == "__main__":
    print("Autonomous Builder for Prototype Generation Platform")
    print("===================================================")
    
    builder = AutonomousBuilder()
    builder.execute_full_setup()
